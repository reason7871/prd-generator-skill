#!/usr/bin/env python3
"""
Convert Markdown PRD to Word document (.docx)
Supports: headings, tables, code blocks, images, lists
"""

import argparse
import re
import os
from pathlib import Path
from typing import Optional, Tuple


def convert_md_to_docx(input_path: str, output_path: Optional[str] = None) -> str:
    """Convert a Markdown file to Word document."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        print("Error: Install python-docx: pip install python-docx")
        return ""

    input_file = Path(input_path)
    if not input_file.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")

    output_path = output_path or str(input_file.with_suffix('.docx'))

    # Read markdown content
    with open(input_file, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Create document
    doc = Document()

    # Configure styles
    _setup_styles(doc)

    # Parse and convert
    lines = md_content.split('\n')
    i = 0
    in_code_block = False
    code_block_content = []
    code_language = ''
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]

        # Code blocks
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_language = line.strip()[3:].strip()
                code_block_content = []
            else:
                in_code_block = False
                _add_code_block(doc, '\n'.join(code_block_content), code_language)
                code_block_content = []
            i += 1
            continue

        if in_code_block:
            code_block_content.append(line)
            i += 1
            continue

        # Tables
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            # Parse table row
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells and not all(c.replace('-', '').replace(':', '') == '' for c in cells):
                table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            in_table = False
            if table_rows:
                _add_table(doc, table_rows)
            table_rows = []

        # Images
        img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line)
        if img_match:
            alt_text, img_path = img_match.groups()
            _add_image(doc, img_path, alt_text, input_file.parent)
            i += 1
            continue

        # Headings
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            _add_heading(doc, text, level)
            i += 1
            continue

        # Horizontal rules
        if re.match(r'^-{3,}$|^\*{3,}$|^_{3,}$', line.strip()):
            doc.add_paragraph('_' * 50)
            i += 1
            continue

        # Blockquotes
        if line.strip().startswith('>'):
            text = line.strip()[1:].strip()
            _add_blockquote(doc, text)
            i += 1
            continue

        # Lists
        list_match = re.match(r'^(\s*)[-*+]\s+(.+)$', line)
        numbered_match = re.match(r'^(\s*)(\d+)\.\s+(.+)$', line)
        if list_match:
            _add_list_item(doc, list_match.group(2), bullet=True)
            i += 1
            continue
        if numbered_match:
            _add_list_item(doc, numbered_match.group(3), bullet=False)
            i += 1
            continue

        # Regular paragraphs
        if line.strip():
            _add_paragraph(doc, line)
        else:
            # Empty line - add small break
            pass

        i += 1

    # Handle any remaining table
    if in_table and table_rows:
        _add_table(doc, table_rows)

    # Add metadata
    core_props = doc.core_properties
    core_props.author = 'PRD Generator'
    core_props.title = input_file.stem

    # Save
    doc.save(output_path)
    print(f"✓ Document saved to: {output_path}")
    return output_path


def _setup_styles(doc: Document):
    """Configure document styles."""
    styles = doc.styles

    # Configure heading styles
    for i in range(1, 7):
        try:
            style = styles[f'Heading {i}']
            font = style.font
            font.bold = True
            if i == 1:
                font.size = Pt(24)
                font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
            elif i == 2:
                font.size = Pt(18)
                font.color.rgb = RGBColor(0x16, 0x21, 0x3e)
            elif i == 3:
                font.size = Pt(14)
                font.color.rgb = RGBColor(0x2d, 0x3e, 0x50)
            else:
                font.size = Pt(12)
        except KeyError:
            pass


def _add_heading(doc: Document, text: str, level: int):
    """Add a heading to the document."""
    # Clean up markdown formatting
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)  # Bold
    text = re.sub(r'\*([^*]+)\*', r'\1', text)  # Italic
    text = re.sub(r'`([^`]+)`', r'\1', text)  # Code

    heading_level = min(level, 9)
    doc.add_heading(text, level=heading_level)


def _add_paragraph(doc: Document, text: str):
    """Add a paragraph with inline formatting."""
    # Handle inline formatting
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)  # Bold - simplified
    text = re.sub(r'\*([^*]+)\*', r'\1', text)  # Italic - simplified
    text = re.sub(r'`([^`]+)`', r'\1', text)  # Code - simplified

    para = doc.add_paragraph(text)
    return para


def _add_code_block(doc: Document, code: str, language: str = ''):
    """Add a code block with monospace formatting."""
    # Add language label if present
    if language:
        label = doc.add_paragraph(f'[{language}]')
        label.runs[0].italic = True
        label.runs[0].font.size = Pt(9)

    # Add code in monospace
    code_para = doc.add_paragraph()
    code_para.paragraph_format.left_indent = Inches(0.5)

    for line in code.split('\n'):
        run = code_para.add_run(line + '\n')
        run.font.name = 'Consolas'
        run.font.size = Pt(9)


def _add_table(doc: Document, rows: list):
    """Add a table to the document."""
    if not rows or len(rows) < 1:
        return

    num_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = row.cells[j]
                cell.text = cell_text.strip()
                # Bold header row
                if i == 0:
                    cell.paragraphs[0].runs[0].bold = True


def _add_image(doc: Document, img_path: str, alt_text: str, base_dir: Path):
    """Add an image to the document."""
    # Handle relative paths
    if not os.path.isabs(img_path):
        full_path = base_dir / img_path
    else:
        full_path = Path(img_path)

    if full_path.exists():
        try:
            doc.add_picture(str(full_path), width=Inches(6))
            # Add caption
            caption = doc.add_paragraph(alt_text)
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.runs[0].italic = True
        except Exception as e:
            doc.add_paragraph(f'[Image: {img_path}] - Error: {e}')
    else:
        doc.add_paragraph(f'[Image not found: {img_path}]')


def _add_blockquote(doc: Document, text: str):
    """Add a blockquote style paragraph."""
    para = doc.add_paragraph(text)
    para.paragraph_format.left_indent = Inches(0.5)
    para.runs[0].italic = True


def _add_list_item(doc: Document, text: str, bullet: bool = True):
    """Add a list item."""
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'\*([^*]+)\*', r'\1', text)

    if bullet:
        doc.add_paragraph(text, style='List Bullet')
    else:
        doc.add_paragraph(text, style='List Number')


def create_docx_from_template(template_vars: dict, output_path: str) -> str:
    """Create a Word document from template variables (alternative method)."""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        print("Error: Install python-docx: pip install python-docx")
        return ""

    doc = Document()

    # Title
    title = doc.add_heading(template_vars.get('project_name', 'Product Requirements Document'), 0)
    title.alignment = 1  # Center

    # Metadata
    meta = doc.add_paragraph()
    meta.add_run(f"Generated: {template_vars.get('date', 'N/A')}\n").italic = True
    meta.add_run(f"Version: {template_vars.get('version', '1.0')}\n").italic = True
    meta.add_run(f"Source: {template_vars.get('source', 'N/A')}").italic = True

    doc.add_paragraph()  # Spacer

    # Sections
    for section in template_vars.get('sections', []):
        doc.add_heading(section['title'], level=1)
        doc.add_paragraph(section['content'])

    doc.save(output_path)
    return output_path


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Convert Markdown PRD to Word document')
    parser.add_argument('input', help='Input Markdown file path')
    parser.add_argument('-o', '--output', help='Output Word file path (optional)')
    args = parser.parse_args()

    result = convert_md_to_docx(args.input, args.output)
    if result:
        print(f"Conversion complete!")
